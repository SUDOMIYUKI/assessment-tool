"""
手書き用支援シート（Word形式）生成スクリプト - 改善版
- 全項目を1ページA4サイズに凝縮
- OCR読み取りに最適化
- チェックボックスは ■ を使用
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """セルの余白を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)


def create_handwriting_sheet(output_path='output/手書き支援シート_1ページ版.docx'):
    """
    手書き用支援シート（1ページ版）を作成
    
    Args:
        output_path: 出力ファイルのパス
    """
    doc = Document()
    
    # ページ設定（A4、余白を最小化）
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4縦
    section.page_width = Cm(21)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # === タイトル ===
    title = doc.add_heading('不登校支援 初回アセスメント手書きシート', 0)
    title_format = title.paragraph_format
    title_format.space_before = Pt(0)
    title_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    
    # === 基本情報を上部バーに配置 ===
    header_table = doc.add_table(rows=2, cols=4)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # 1行目
    row1_cells = header_table.rows[0].cells
    row1_cells[0].text = '記入日：    /  /  '
    row1_cells[1].text = '児童氏名：            '
    row1_cells[2].text = '保護者氏名：            '
    row1_cells[3].text = '支援員：          '
    
    # 2行目
    row2_cells = header_table.rows[1].cells
    row2_cells[0].text = '学校名：          '
    row2_cells[1].text = '学年：    '
    row2_cells[2].text = '性別：□男性 □女性'
    row2_cells[3].text = 'ひとり親：□該当 □非該当'
    
    for row in header_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_margins(cell, top=20, start=30, bottom=20, end=30)
    
    # 余白調整
    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(3)
    
    # === メインコンテンツ（2カラム表） ===
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    main_table.allow_autofit = False
    
    # 列幅を設定
    main_table.columns[0].width = Cm(9.0)
    main_table.columns[1].width = Cm(9.0)
    
    left_cell = main_table.rows[0].cells[0]
    right_cell = main_table.rows[0].cells[1]
    
    # セルの余白を最小化
    set_cell_margins(left_cell, top=40, start=40, bottom=40, end=40)
    set_cell_margins(right_cell, top=40, start=40, bottom=40, end=40)
    
    # ========== 左カラム ==========
    
    # セクション1：家族関係（ジェノグラム）
    p = left_cell.add_paragraph()
    run = p.add_run('【家族関係・ジェノグラム】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(2)
    
    # ジェノグラム記入欄
    genogram_items = [
        '記号：□男性  ○女性  ■本人',
        '　　　―同居  ･･･別居  ×死亡',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '',
        '特記事項：',
        '',
    ]
    
    for item in genogram_items:
        p = left_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.name = 'MS Gothic'  # 等幅フォントで図形を整列
    
    # セクション2：登校状況
    p = left_cell.add_paragraph()
    run = p.add_run('【登校状況】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    
    attendance_items = [
        '頻度：□週0  □週1-2  □週3-4  □毎日',
        '不登校：□該当  □非該当',
        '詳細：',
        '',
    ]
    
    for item in attendance_items:
        p = left_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # セクション3：生活状況
    p = left_cell.add_paragraph()
    run = p.add_run('【生活状況】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    life_items = [
        'リズム：□朝起きられない □昼夜逆転',
        '　　　　□睡眠不足 □問題なし',
        '習慣：□食事の乱れ □運動不足',
        '　　　□ゲーム依存 □問題なし',
        '外出：□する □コンビニ程度 □しない',
        '詳細：',
        '',
    ]
    
    for item in life_items:
        p = left_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # セクション4：学習状況
    p = left_cell.add_paragraph()
    run = p.add_run('【学習状況】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    study_items = [
        '□学習の遅れ □低学力 □習慣なし',
        '□環境なし □問題なし',
        '詳細：',
        '',
    ]
    
    for item in study_items:
        p = left_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # セクション5：対人関係
    p = left_cell.add_paragraph()
    run = p.add_run('【対人関係】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    social_items = [
        '□対人緊張 □友達不安 □コミュ苦手',
        '□問題なし',
        '詳細：',
        '',
    ]
    
    for item in social_items:
        p = left_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # ========== 右カラム ==========
    
    # セクション4：発達特性・医療情報
    p = right_cell.add_paragraph()
    run = p.add_run('【発達特性・医療情報】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(2)
    
    medical_items = [
        '□発達特性あり（内容：        ）',
        '□通院（病院：        頻度：   ）',
        '□診断（名称：                ）',
        '□投薬（薬名：                ）',
        '□手帳（種類：                ）',
    ]
    
    for item in medical_items:
        p = right_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # セクション5：家庭環境
    p = right_cell.add_paragraph()
    run = p.add_run('【家庭環境】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    family_items = [
        '課題：□経済困難 □家族関係',
        '　　　□他世帯員 □虐待 □その他',
        '詳細：',
        '',
    ]
    
    for item in family_items:
        p = right_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # セクション6：短期目標
    p = right_cell.add_paragraph()
    run = p.add_run('【短期目標】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    short_items = [
        '課題：',
        '',
        '現状：',
        '',
        'ニーズ本人：',
        '',
        'ニーズ保護者：',
        '',
        '目標：',
        '',
        '方法：',
        '',
    ]
    
    for item in short_items:
        p = right_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # セクション7：長期目標
    p = right_cell.add_paragraph()
    run = p.add_run('【長期目標（本事業）】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    long_items = [
        '課題：',
        '',
        '目標：',
        '',
        '方法：',
        '',
    ]
    
    for item in long_items:
        p = right_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # セクション8：支援への希望
    p = right_cell.add_paragraph()
    run = p.add_run('【支援への希望】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    support_items = [
        '曜日：□月 □火 □水 □木 □金',
        '時間：',
        '場所：',
        '支援員希望：',
        '',
        '解決したいこと：',
        '',
        '',
    ]
    
    for item in support_items:
        p = right_cell.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # === フッター（当日の様子） ===
    doc.add_paragraph()  # 空行
    
    p = doc.add_paragraph()
    run = p.add_run('【当日の様子・メモ】')
    run.font.size = Pt(10)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    
    memo_lines = ['' for _ in range(5)]
    for line in memo_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8)
    
    # ファイルを保存
    doc.save(output_path)
    print(f'✅ 手書き支援シート（1ページ版）を作成しました: {output_path}')
    
    return output_path


if __name__ == '__main__':
    from pathlib import Path
    
    # 出力ディレクトリを作成
    output_dir = Path('output')
    output_dir.mkdir(exist_ok=True)
    
    output_path = output_dir / '手書き支援シート_1ページ版.docx'
    create_handwriting_sheet(str(output_path))

